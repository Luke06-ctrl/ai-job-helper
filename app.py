"""职途星 — AI 大学生就业助手"""

import base64
import datetime
import io
import json
import os
import uuid

from flask import Flask, render_template, request, jsonify, Response, session, send_file
from flask_limiter import Limiter
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import SECRET_KEY, DEBUG
from utils.deepseek import chat, chat_stream, ocr_image
from prompts.resume import build_messages as resume_build
from prompts.interview import (
    build_start_message,
    build_evaluate_message,
    INTERVIEW_TYPES,
)
from prompts.career import build_messages as career_build, PRESET_QUESTIONS

app = Flask(__name__)
app.secret_key = SECRET_KEY


# 获取真实客户端 IP（穿透代理/负载均衡）
def get_real_ip():
    forwarded = request.headers.get("X-Forwarded-For", "")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.remote_addr or "127.0.0.1"


# 请求频率限制（防刷 API 额度）
limiter = Limiter(
    get_real_ip,
    app=app,
    default_limits=["30 per minute"],  # 全局默认
    storage_uri="memory://",
)


@app.errorhandler(429)
def ratelimit_error(e):
    return jsonify({"error": "请求太频繁，请稍后再试"}), 429


# ═══════════════ 页面路由 ═══════════════

@app.route("/")
def index():
    """首页"""
    return render_template("index.html")


@app.route("/resume")
def resume_page():
    """简历优化页"""
    return render_template("resume.html")


@app.route("/interview")
def interview_page():
    """AI 模拟面试页"""
    return render_template("interview.html", types=INTERVIEW_TYPES)


@app.route("/career")
def career_page():
    """职业规划问答页"""
    return render_template("career.html", presets=PRESET_QUESTIONS)


# ═══════════════ API 路由 ═══════════════

@app.route("/api/resume/analyze", methods=["POST"])
@limiter.limit("3 per minute")
def api_resume_analyze():
    """简历优化 API"""
    data = request.get_json()
    resume_text = data.get("content", "").strip()
    target_job = data.get("target_job", "").strip()

    if not resume_text:
        return jsonify({"error": "请输入简历内容"}), 400
    if len(resume_text) < 50:
        return jsonify({"error": "简历内容太短，请至少输入 50 个字"}), 400

    messages = resume_build(resume_text, target_job)
    result = chat(messages, temperature=0.6)
    return jsonify({"result": result})


@app.route("/api/resume/upload", methods=["POST"])
@limiter.limit("3 per minute")
def api_resume_upload():
    """文件上传 + 图片 OCR 提取"""
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "请选择文件"}), 400

    filename = file.filename.lower()
    content = ""

    try:
        # 图片文件 → OCR 识别
        if filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
            raw = file.read()
            img_b64 = base64.b64encode(raw).decode("utf-8")
            mime = "image/" + ("jpeg" if filename.endswith(".jpg") or filename.endswith(".jpeg") else filename.split(".")[-1])
            if filename.endswith(".webp"):
                mime = "image/webp"
            elif filename.endswith(".bmp"):
                mime = "image/bmp"
            content = ocr_image(img_b64, mime)
            if content.startswith("[OCR错误]"):
                return jsonify({"error": content}), 500

        # PDF 文件
        elif filename.endswith(".pdf"):
            raw = file.read()
            reader = PdfReader(io.BytesIO(raw))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
            if not content.strip():
                return jsonify({"error": "未能从 PDF 中提取文字，请尝试粘贴或拍照"}), 400

        # Word 文件
        elif filename.endswith(".docx"):
            raw = file.read()
            doc = Document(io.BytesIO(raw))
            content = "\n".join(p.text for p in doc.paragraphs)
            if not content.strip():
                return jsonify({"error": "Word 文件内容为空"}), 400

        # 纯文本文件
        elif filename.endswith(".txt"):
            raw = file.read()
            content = raw.decode("utf-8", errors="replace")

        else:
            return jsonify({"error": "不支持的文件格式，请上传 PDF/Word/TXT/图片"}), 400

        if not content.strip() or len(content.strip()) < 20:
            return jsonify({"error": "提取的文字内容太少，请确认文件包含简历内容"}), 400

        return jsonify({"content": content.strip(), "filename": file.filename})

    except Exception as e:
        return jsonify({"error": f"文件处理失败: {str(e)}"}), 500


@app.route("/api/interview/start", methods=["POST"])
@limiter.limit("3 per minute")
def api_interview_start():
    """面试开始 API"""
    data = request.get_json()
    job_type = data.get("job_type", "general")

    if job_type not in INTERVIEW_TYPES:
        return jsonify({"error": "无效的岗位类型"}), 400

    # 初始化会话状态
    session["interview_job"] = job_type
    session["interview_round"] = 1
    session["interview_history"] = []

    messages = build_start_message(job_type, round_num=1)
    result = chat(messages, temperature=0.8)

    # 保存 AI 的问题到历史
    session["interview_history"].append({"role": "ai", "content": result})
    session.modified = True

    return jsonify({
        "result": result,
        "round": 1,
        "job_name": INTERVIEW_TYPES[job_type]["name"],
    })


@app.route("/api/interview/answer", methods=["POST"])
@limiter.limit("10 per minute")
def api_interview_answer():
    """面试回答 API"""
    data = request.get_json()
    answer = data.get("answer", "").strip()

    job_type = session.get("interview_job", "general")
    round_num = session.get("interview_round", 1)
    history = session.get("interview_history", [])

    if not answer:
        return jsonify({"error": "请输入你的回答"}), 400

    # 保存用户回答
    history.append({"role": "user", "content": answer})

    if answer == "结束面试" or round_num >= 6:
        # 生成评价报告
        history_text = "\n".join([
            f"面试官: {h['content']}" if h["role"] == "ai"
            else f"候选人: {h['content']}"
            for h in history
        ])
        messages = build_evaluate_message(history_text)
        result = chat(messages, temperature=0.6)

        # 清理会话
        session.pop("interview_job", None)
        session.pop("interview_round", None)
        session.pop("interview_history", None)

        return jsonify({
            "result": result,
            "round": round_num,
            "finished": True,
        })

    # 继续面试
    round_num += 1
    session["interview_round"] = round_num

    # 构建上下文理解
    recent = history[-4:]  # 最近 2 轮
    context = "\n".join([
        f"{'面试官' if h['role'] == 'ai' else '候选人'}: {h['content']}"
        for h in recent
    ])

    followup_messages = [
        {
            "role": "system",
            "content": f"""你是一位经验丰富的面试官。

岗位类型：{INTERVIEW_TYPES[job_type]['name']}
当前是第 {round_num} 轮问答。

## 面试历史
{context}

请根据候选人上一轮的回答，提出第 {round_num} 个问题。规则：
1. 针对候选人回答中的细节进行追问
2. 只问一个问题
3. 如果这是第 5-6 轮，可以开始收尾，问候选人是否有问题要问你""",
        },
        {"role": "user", "content": "请提出下一个面试问题。"},
    ]

    result = chat(followup_messages, temperature=0.8)
    history.append({"role": "ai", "content": result})
    session["interview_history"] = history
    session.modified = True

    return jsonify({
        "result": result,
        "round": round_num,
        "finished": False,
    })


@app.route("/api/career/chat", methods=["POST"])
@limiter.limit("5 per minute")
def api_career_chat():
    """职业规划问答 API"""
    data = request.get_json()
    question = data.get("question", "").strip()
    history = data.get("history", [])

    if not question:
        return jsonify({"error": "请输入你的问题"}), 400

    messages = career_build(question, history)
    result = chat(messages, temperature=0.8, max_tokens=2048)
    return jsonify({"result": result})


@app.route("/api/career/chat/stream", methods=["POST"])
@limiter.limit("5 per minute")
def api_career_chat_stream():
    """职业规划问答 API（流式）"""
    data = request.get_json()
    question = data.get("question", "").strip()
    history = data.get("history", [])

    if not question:
        return jsonify({"error": "请输入你的问题"}), 400

    messages = career_build(question, history)

    def generate():
        for chunk in chat_stream(messages, temperature=0.8, max_tokens=2048):
            yield f"data: {json.dumps({'content': chunk})}\n\n"
        yield "data: [DONE]\n\n"

    return Response(generate(), mimetype="text/event-stream")


# ═══════════════ 导出 Word 简历 ═══════════════

@app.route("/api/resume/export", methods=["POST"])
@limiter.limit("3 per minute")
def api_resume_export():
    """将优化内容导出为 Word 文档"""
    data = request.get_json()
    resume_text = data.get("content", "").strip()
    target_job = data.get("target_job", "").strip()

    if not resume_text:
        return jsonify({"error": "请输入简历内容"}), 400

    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_margin_top = Cm(2)
    section.page_margin_bottom = Cm(2)
    section.page_margin_left = Cm(2.5)
    section.page_margin_right = Cm(2.5)

    # 标题
    title = doc.add_heading("个人简历", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if target_job:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"目标岗位：{target_job}")
        run.font.size = Pt(11)
        run.font.color.rgb = None

    doc.add_paragraph()  # 空行

    # 解析简历内容并结构化
    lines = resume_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 检测标题行（包含冒号或关键词）
        if "：" in line or ":" in line or any(kw in line for kw in ["信息", "教育", "经历", "实习", "项目", "技能", "证书", "荣誉", "自我"]):
            p = doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(10.5)

    # 页脚
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("— 本简历由 职途星 AI 助手优化生成 —")
    run.font.size = Pt(8)
    run.font.color.rgb = None

    # 返回 Word 文件
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = "简历优化_职途星.docx"
    if target_job:
        filename = f"简历_{target_job}_职途星.docx"

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ═══════════════ 面试历史记录 ═══════════════

HISTORY_DIR = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(HISTORY_DIR, exist_ok=True)


def _save_history(record):
    """保存面试记录到 JSON 文件"""
    rid = uuid.uuid4().hex[:8]
    record["id"] = rid
    record["time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    path = os.path.join(HISTORY_DIR, f"{rid}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(record, f, ensure_ascii=False, indent=2)
    return rid


def _load_all_history():
    """加载所有面试历史"""
    records = []
    for fname in os.listdir(HISTORY_DIR):
        if fname.endswith(".json"):
            path = os.path.join(HISTORY_DIR, fname)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    records.append(json.load(f))
            except Exception:
                pass
    records.sort(key=lambda r: r.get("time", ""), reverse=True)
    return records


@app.route("/history")
def history_page():
    """面试历史页"""
    records = _load_all_history()
    return render_template("history.html", records=records, types=INTERVIEW_TYPES)


@app.route("/api/interview/save", methods=["POST"])
@limiter.limit("5 per minute")
def api_interview_save():
    """保存面试结果"""
    data = request.get_json()
    job_type = data.get("job_type", "general")
    report = data.get("report", "")
    feedback = data.get("feedback", "")

    if not report:
        return jsonify({"error": "缺少面试报告"}), 400

    record = {
        "job_type": job_type,
        "job_name": INTERVIEW_TYPES.get(job_type, {}).get("name", job_type),
        "report": report,
        "feedback": feedback,
    }
    rid = _save_history(record)
    return jsonify({"id": rid, "ok": True})


@app.route("/api/history/list", methods=["GET"])
def api_history_list():
    """获取历史记录列表"""
    return jsonify(_load_all_history())


@app.route("/api/history/<hid>", methods=["GET"])
def api_history_detail(hid):
    """获取单条面试历史"""
    path = os.path.join(HISTORY_DIR, f"{hid}.json")
    if not os.path.exists(path):
        return jsonify({"error": "记录不存在"}), 404
    with open(path, "r", encoding="utf-8") as f:
        return jsonify(json.load(f))


# ═══════════════ 启动入口 ═══════════════

if __name__ == "__main__":
    app.run(debug=DEBUG, host="0.0.0.0", port=5000)
